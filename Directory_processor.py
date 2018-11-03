
# coding: utf-8

# In[86]:


import os
import nltk
import sys
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.style import WD_STYLE_TYPE
from operator import itemgetter
from nltk.tag import StanfordNERTagger
from nltk.tokenize import word_tokenize
try:
    from xml.etree.cElementTree import XML
except ImportError:
    from xml.etree.ElementTree import XML
import zipfile
import re
from datetime import datetime,date


# In[57]:


st = StanfordNERTagger('/home/sidhantj/Documents/Notebooks/stanford/english.all.3class.distsim.crf.ser.gz',
                       '/home/sidhantj/Documents/Notebooks/stanford/stanford-ner.jar',
                       encoding='utf-8') 


# In[58]:


directory = os.listdir('/home/sidhantj/Documents/Notebooks/Processed_Folder')
os.chdir('/home/sidhantj/Documents/Notebooks/Processed_Folder')
titleKw = {"Article Title", "Title", "Paper Title", "titled", "Titled"}
abstractKw = {'Abstract', 'Summary', 'Precis', 'Clinical significance', 'Snapshot', 'Capsule', 'Overview'}
introductionKw = {'Introduction', 'Literature review', 'Context', 'Background'}
methodKw = {'Methods', 'Case presentation', 'Case report', 'Case', 'Case Description', 'Presentation of the Case',
            'Diagnosis and Treatment', 'Study design', 'Materials and Methods', 'Apparatus', 'Methodology',
            'Experimental'}
keykw = {'Keyword', 'Keywords', 'keywords', 'key words', 'Key Words', 'keyword', 'Key words', 'KeyWords'}
resultKw = {'Result', 'Results', 'Discussion', 'Discussions', 'Results and Discussions'}


# In[59]:


# -- coding: UTF-8 --
# def get_human_names(testext):
#     tokens = nltk.tokenize.word_tokenize(testext)
#     pos = nltk.pos_tag(tokens)
#     sentt = nltk.ne_chunk(pos, binary = False)
#     person_list = []
#     person = []
#     name = ""
#     for subtree in sentt.subtrees(filter=lambda t: t.label() == 'PERSON'):
#         for leaf in subtree.leaves():
#             person.append(leaf[0])
#         if len(person) > 1: #avoid grabbing lone surnames
#             for part in person:
#                 name += part + ' '
#             if name[:-1] not in person_list:
#                 person_list.append(name[:-1])
#             name = ''
#         person = []

#     return (person_list)


# In[88]:


#highlighting run, styling para containng author
def get_names(paragraphs,document):
    testext = ""
    for paragraph in paragraphs[1:7]:
        testext += " \n"
        testext += (paragraph.text)
    
    tokenized_text = word_tokenize(testext)
    classified_text = st.tag(tokenized_text)
    person_list =[item[0] for item in classified_text if item[1] == 'PERSON']
    result = False
    for name in person_list:
        for paragraph in document.paragraphs[1:7]:
            if len(list(word_tokenize(paragraph.text))) <= 60:
                if name in paragraph.text:
                    paragraph.style = 'Author'
                    #for run in paragraph.runs:
                     #   if name in run.text:
                    result = True
                            #run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    document.save(file.split(".")[0]+"_PROCESSED.docx")
    #print (names)
    return result
    


# In[61]:


def get_remaining_part(list_of_words, split_by):
    if (list_of_words[0] == split_by and len(list_of_words) == 1):
        #print ("No other words in sentence")
        return (split_by)
    else:    
        try:
            next_word = list_of_words[list_of_words.index(split_by) + 1:len(list_of_words)]
        except ValueError:
            next_word = "no word found"
        #print(next_word)
        return next_word


# In[62]:


def keyword_score(paragraph):
    result = 0
    sent_text = nltk.sent_tokenize(paragraph.text)
    for sent in sent_text:
        tokenized_text = nltk.word_tokenize(sent)
        word_list = list(filter(lambda word: word not in ',.', tokenized_text))
        #print (word_list)
        if len(word_list) <= 25:
                result = 1
                for word in word_list:
                    if word in titleKw:
                        #print ('Contains keyword:  ', sent_text)
                        result = result + 1
                                #check whether title in same line of Keyword or next line and get article paragraph 
                        possible_title = get_remaining_part(tokenized_text, word)
                                # further check for same line or next line can be title
                        #print ('Possible title:', possible_title, '\nOther text in the para: ', paragraph.text)
    return result


# In[63]:


def is_font_bold(paragraph):
    result = 0
    if paragraph:
        #print ('paragraph:', paragraph)
        for run in paragraph.runs:
            if run.bold is not None:
                #print (run.text)
                result = 1
    return result


# In[64]:


def max_val(l, i):
    
    return max(enumerate(map(itemgetter(i), l)), key=itemgetter(1))


# In[65]:


def delete_paragraph(_paragraph):
    p = _paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


# In[66]:


def get_title_edit(paragraphs,document):
    result=False
    paraNoResult = []
    index = 0
    for paragraph in paragraphs[0:5]:
        resultOfScore = {"keywordScore": keyword_score(paragraph),
                         "boldScore": is_font_bold(paragraph)}
        paraNoResult.append((resultOfScore, index))
    index += 1

    #print(paraNoResult)
    paraNo = 0
    aggregate_score_result = []
    for result in paraNoResult:
        #print ("para number = ", paraNo, 'score = ', result[0]['keywordScore'] + result[0]['boldScore'])
        aggregate_score_result.append([paraNo, result[0]['keywordScore'] + result[0]['boldScore']])
        paraNo = paraNo + 1
    #print ('aggregate score result = ', aggregate_score_result)
        paragraphs[max_val(aggregate_score_result, -1)[0]].style = 'Title'
        #print (paragraphs[max_val(aggregate_score_result, -1)[0]].text)
        #for run in paragraphs[max_val(aggregate_score_result, -1)[0]].runs:
            #run.font.highlight_color = WD_COLOR_INDEX.RED
        document.save(file.split(".")[0]+"_PROCESSED.docx")
        result = True
    return result


# In[67]:


def get_abstract_edit(paragraphs,document):
    result=False
    index = 0
    for paragraph in paragraphs:
        sent_text = nltk.sent_tokenize(paragraph.text)  # this gives us a list of sentences
        for sentence in sent_text:
            tokenized_text = nltk.word_tokenize(sentence)
            if len(list(filter(lambda _word: _word not in ',.', tokenized_text))) <= 2:
                for word in list(filter(lambda _word: _word not in ',.', tokenized_text)):
                    if word in abstractKw:
                        paragraph.style = 'Abstract'
                        if len(paragraph.text) <= 25:
                            paragraphs[index+1].style = 'Abstract'
                        document.save(file.split(".")[0]+"_PROCESSED.docx")
                        result = True
        index += 1
    return result


# In[68]:


def get_intro_edit(paragraphs,document):
    result=False
    index = 0
    for paragraph in paragraphs:
        sent_text = nltk.sent_tokenize(paragraph.text)  # this gives us a list of sentences
        for sentence in sent_text:
            tokenized_text = nltk.word_tokenize(sentence)
            if len(list(filter(lambda _word: _word not in ',.', tokenized_text))) <= 2:
                for word in list(filter(lambda _word: _word not in ',.', tokenized_text)):
                    if word in introductionKw:
                        #print ('Contains instroduction keyword:', paragraph.text)
                        paragraph.style = 'Introduction'
                        if len(paragraph.text) <= 25:
                            paragraphs[index+1].style = 'Introduction'
                        document.save(file.split(".")[0]+"_PROCESSED.docx")
                        result = True
        index += 1
    return result


# In[69]:


def get_method_edit(paragraphs,document):
    result=False
    index = 0
    for paragraph in paragraphs:
        sent_text = nltk.sent_tokenize(paragraph.text)  # this gives us a list of sentences
        for sentence in sent_text:
            tokenized_text = nltk.word_tokenize(sentence)
            if len(list(filter(lambda _word: _word not in ',.', tokenized_text))) <= 3:
                for word in list(filter(lambda _word: _word not in ',.', tokenized_text)):
                    if word in methodKw:
                        #print ('Contains instroduction keyword:', paragraph.text)
                        paragraph.style = 'Method'
                        if len(paragraph.text) <= 25:
                            paragraphs[index+1].style = 'Method'
                        document.save(file.split(".")[0]+"_PROCESSED.docx")
                        result = True
        index += 1
    return result


# In[70]:


def get_result_edit(paragraphs,document):
    result=False
    index = 0
    for paragraph in paragraphs:
        sent_text = nltk.sent_tokenize(paragraph.text)  # this gives us a list of sentences
        for sentence in sent_text:
            tokenized_text = nltk.word_tokenize(sentence)
            if len(list(filter(lambda _word: _word not in ',.', tokenized_text))) <= 3:
                for word in list(filter(lambda _word: _word not in ',.', tokenized_text)):
                    if word in resultKw:
                        #print ('Contains instroduction keyword:', paragraph.text)
                        paragraph.style = 'Result'
                        if len(paragraph.text) <= 25:
                            paragraphs[index+1].style = 'Result'
                        document.save(file.split(".")[0]+"_PROCESSED.docx")
                        result = True
        index += 1
    return result


# In[71]:


def get_key(paragraphs,document):
    result = False
    index = 0
    for paragraph in paragraphs:
        sent_text = nltk.sent_tokenize(paragraph.text)  # this gives us a list of sentences
        for sentence in sent_text:
            tokenized_text = nltk.word_tokenize(sentence)
            #if len(list(filter(lambda _word: _word not in ',.', tokenized_text))) <= 3:
             #    for word in list(filter(lambda _word: _word not in ',.', tokenized_text)):
              #       print (word)
            #print (list(tokenized_text))
            windex = 0
            for word in list(tokenized_text):
                if word.lower() == 'key' and (tokenized_text[list(tokenized_text).index(word)+1].lower() == 'words' or tokenized_text[list(tokenized_text).index(word)+1].lower() == 'word'):
                    #print (word.lower(),tokenized_text[list(tokenized_text).index(word)+1].lower())
                    #if tokenized_text[list(tokenized_text).index(word)+1].lower() == 'word' or tokenized_text[list(tokenized_text).index(word)+1].lower() == 'words':
                    paragraph.style = 'Keywords'
                    if len(paragraph.text) <= 15:
                        paragraphs[index+1].style = 'Keywords'
                    document.save(file.split(".")[0]+"_PROCESSED.docx")
                    result = True
                elif word in keykw:
                    paragraph.style = 'Keywords'
                    if len(paragraph.text) <= 15:
                        paragraphs[index+1].style = 'Keywords'
                    document.save(file.split(".")[0]+"_PROCESSED.docx")
                    result = True
                windex += 1
        index += 1
    return result


# In[72]:


def find_cities(paragraphs,document):
    result = False
    text = ''
    for paragraph in paragraphs[1:]:
        text += paragraph.text
        token_text = word_tokenize(text)
        classified_text = st.tag(token_text)
        location_list =[item[0] for item in classified_text if item[1] == 'LOCATION']
        organization_list = [item[0] for item in classified_text if item[1] == 'ORGANIZATION']
    for paragraph in paragraphs[1:]:
        sent_text = nltk.sent_tokenize(paragraph.text)  # this gives us a list of sentences
        for sentence in sent_text:
            tokenized_text = nltk.word_tokenize(sentence)
            for word in list(tokenized_text):
                #print (word)
                if word in list(location_list):
                    if word in list(organization_list):
                        paragraph.style = 'Affiliation'
                        document.save(file.split(".")[0]+"_PROCESSED.docx")
                        result = True
    #print (location_list)
    return result


# In[73]:


# def process_file(directory):
#     for file in directory:
#         print (file)
#         output = {"filename" : file, "authornames" : False, "titlename" : False, "abstract" : False, "intro" : False,
#          "method" : False, 'keywords': False, 'affi':False}
#         document = Document(file)
#         paragraphs = list(document.paragraphs)
#         styles = document.styles
#         allStyles = []
#         paragraph_styles = [s for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH]
#         for style12 in paragraph_styles:
#             allStyles.append(style12.name)
#         #print (allStyles)
#         if 'Author' not in allStyles:
#             style = styles.add_style('Author', WD_STYLE_TYPE.PARAGRAPH)
#         if 'Title' not in allStyles:
#             style = styles.add_style('Title', WD_STYLE_TYPE.PARAGRAPH)
#         if 'Abstract' not in allStyles:
#             style = styles.add_style('Abstract', WD_STYLE_TYPE.PARAGRAPH) 
#         if 'Introduction' not in allStyles:
#             style = styles.add_style('Introduction', WD_STYLE_TYPE.PARAGRAPH)
#         if 'Method' not in allStyles:
#             style = styles.add_style('Method', WD_STYLE_TYPE.PARAGRAPH)
#         if 'Keywords' not in allStyles:
#             style = styles.add_style('Keywords', WD_STYLE_TYPE.PARAGRAPH)
#         if 'Affiliation' not in allStyles:
#             style = styles.add_style('Affiliation', WD_STYLE_TYPE.PARAGRAPH)
#         document.save(file)
#         paragraphs = list(document.paragraphs)
#         output["authornames"] = get_names(paragraphs)
#         output["titlename"] = get_title_edit(paragraphs)
#         output["abstract"] = get_abstract_edit(paragraphs)
#         output["intro"] = get_intro_edit(paragraphs)
#         output["method"] = get_method_edit(paragraphs)
#         output['keywords'] = get_key(paragraphs)
#         output['affi'] = find_cities(paragraphs)
#         print(output, " \n\n")


# In[74]:


#process_file(directory)


# In[75]:


def get_email_para(path):
    document = zipfile.ZipFile(path)
    xml_content = document.read('word/document.xml')
    document.close()
    tree = XML(xml_content)
    WORD_NAMESPACE = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    PARA = WORD_NAMESPACE + 'p'
    TEXT = WORD_NAMESPACE + 't'
    paragraphs = []
    for paragraph in tree.getiterator(PARA):
        texts = [node.text
        for node in paragraph.getiterator(TEXT)
            if node.text]
        if texts:
            paragraphs.append(''.join(texts))
    text = ''
    for paragraph in paragraphs:
        text += ' '+paragraph
    #print (text)
    result = ''
    reresult = False
    match = re.findall(r'[\w\.-]+@[\w\.-]+', text)
    for paragraph in paragraphs:
        if match[0] in paragraph:
            result = (paragraph.split(match[0])[0])
            #print (result)
    document = Document(path)
    paras = list(document.paragraphs)
    for paragraph in paras:
        if result in paragraph.text:
            paragraph.style = 'Correspondence'
            document.save(file.split(".")[0]+"_PROCESSED.docx")
            reresult = True
    return reresult


# In[76]:


def get_references(paragraphs,document):
    result = False
    index = 0
    p = r"[0-9]{4}"
    for paragraph in paragraphs:
        sent_text = nltk.sent_tokenize(paragraph.text) 
        for sentence in sent_text:
            tokenized_text = nltk.word_tokenize(sentence)
            for word in list(tokenized_text):
                if (word.lower() == 'references' or word.lower() == 'reference') and len(tokenized_text) <= 2:
                    paragraph.style = 'References'
                    result = True
                    if len(paragraph.text) <= 15:
                        paragraphs[index+1].style = 'References'
                document.save(file.split(".")[0]+"_PROCESSED.docx")
        tokenized_text = nltk.word_tokenize(paragraph.text)
        classified_text = st.tag(tokenized_text)
        names_list =  [item[0] for item in classified_text if item[1] == 'PERSON']
        year_list = re.findall(p,paragraph.text)
        if len(names_list) != 0:
            if len(year_list) != 0:
                if len(tokenized_text) <= 60:
                    paragraph.style = 'References'
                    document.save(file.split(".")[0]+"_PROCESSED.docx")
                    result = True
        index+=1
    return result


# In[89]:


def process_file(file):
    document = Document(file)
    paragraphs = list(document.paragraphs)
    nonemptyparas = []
    for para in paragraphs:
        if para.text != '':
            nonemptyparas.append(para)
    tenparas = nonemptyparas[0:10]
    styles = document.styles
    allStyles = []
    paragraph_styles = [s for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH]
    for style in paragraph_styles:
        allStyles.append(style.name)
    #print (allStyles, " \n\n")
    if 'Author' not in allStyles:
        style = styles.add_style('Author', WD_STYLE_TYPE.PARAGRAPH)
    if 'Title' not in allStyles:
        style = styles.add_style('Title', WD_STYLE_TYPE.PARAGRAPH)
    if 'Abstract' not in allStyles:
        style = styles.add_style('Abstract', WD_STYLE_TYPE.PARAGRAPH) 
    if 'Introduction' not in allStyles:
        style = styles.add_style('Introduction', WD_STYLE_TYPE.PARAGRAPH)
    if 'Method' not in allStyles:
        style = styles.add_style('Method', WD_STYLE_TYPE.PARAGRAPH)
    if 'Keywords' not in allStyles:
        style = styles.add_style('Keywords', WD_STYLE_TYPE.PARAGRAPH)
    if 'Affiliation' not in allStyles:
        style = styles.add_style('Affiliation', WD_STYLE_TYPE.PARAGRAPH)
    if 'Correspondence' not in allStyles:
        style = styles.add_style('Correspondence', WD_STYLE_TYPE.PARAGRAPH)
    if 'Result' not in allStyles:
        style = styles.add_style('Result', WD_STYLE_TYPE.PARAGRAPH)
    if 'References' not in allStyles:
        style = styles.add_style('References', WD_STYLE_TYPE.PARAGRAPH)
    document.save(file)
    output = {"filename" : file, "authornames" : False, "titlename" : False, "abstract" : False, "intro" : False,
             "method" : False, 'keywords': False, 'affi':False, 'corres' : False, 'result' : False, 'reff': False}
    output['reff'] = get_references(paragraphs,document)
    output["authornames"] = get_names(tenparas,document)
    output["titlename"] = get_title_edit(tenparas,document)
    output["abstract"] = get_abstract_edit(paragraphs,document)
    output["intro"] = get_intro_edit(paragraphs,document)
    output["method"] = get_method_edit(paragraphs,document)
    output['keywords'] = get_key(paragraphs,document)
    output['affi'] = find_cities(tenparas,document)
#     output['corres'] = get_email_para(file)
    output['result'] = get_result_edit(paragraphs,document)

    print(output, " \n")


# In[87]:


file_list = []
for file in directory:
    file_list.append(file)
print(file_list,"\n")
for file in file_list:
    time_started = datetime.now().time()
    print ("Work started on file:",file," at:",time_started,"\n")
    process_file(file)
    time_finished = datetime.now().time()
    duration = datetime.combine(date.min, time_finished) - datetime.combine(date.min, time_started)
    print("File took",duration,"\n")

